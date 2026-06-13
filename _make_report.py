# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"C:\Users\이현렬\tuk 2학년1학기\윈도우프로그래밍\Unrailed"
IMG  = os.path.join(BASE, "Unrailed", "Unrailed", "Image")
OUT  = os.path.join(BASE, "최종결과보고서.docx")
GS   = os.path.join(IMG, "scene", "gamestart.png")
MAP  = os.path.join(IMG, "Map", "newUnrailedmap.png")
FONT = "맑은 고딕"

doc = Document()
sec = doc.sections[0]
sec.page_width = Mm(210); sec.page_height = Mm(297)
sec.top_margin = Mm(20); sec.bottom_margin = Mm(18)
sec.left_margin = Mm(22); sec.right_margin = Mm(22)

st = doc.styles['Normal']
st.font.name = FONT; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
st.paragraph_format.line_spacing = 1.25
st.paragraph_format.space_after = Pt(4)

ACCENT = RGBColor(0xC0,0x5A,0x10); DARK = RGBColor(0x1F,0x2A,0x44)

def set_font(run, size=None, bold=None, color=None, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if size: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color: run.font.color.rgb = color

def h1(num, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{num}. {text}"); set_font(r, 14.5, True, DARK)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'8'); bottom.set(qn('w:space'),'4'); bottom.set(qn('w:color'),'C05A10')
    pbdr.append(bottom); pPr.append(pbdr); return p

def h2(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run("■ " + text); set_font(r, 11.5, True, RGBColor(0x33,0x33,0x33)); return p

def body(text, size=10.5, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); set_font(r, size); return p

def bullet(text, size=10.5, bold_head=None):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(2)
    if bold_head:
        r=p.add_run(bold_head); set_font(r,size,True)
        r2=p.add_run(text); set_font(r2,size)
    else:
        r = p.add_run(text); set_font(r, size)
    return p

def caption(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); set_font(r, 9, False, RGBColor(0x66,0x66,0x66)); return p

def add_image(path, width_in=6.3, cap=None):
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width_in))
        if cap: caption(cap)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),hexcolor); tcPr.append(shd)

def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hc = t.rows[0].cells
    for i,htxt in enumerate(headers):
        hc[i].text=''
        r=hc[i].paragraphs[0].add_run(htxt); set_font(r,10,True,RGBColor(0xFF,0xFF,0xFF))
        hc[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; shade(hc[i],'2F3C7E')
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=''; r=cells[i].paragraphs[0].add_run(str(val)); set_font(r,10)
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2); return t

# ===== 표지 =====
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(6)
r=p.add_run("2026 윈도우 프로그래밍 텀 프로젝트"); set_font(r,12,True,RGBColor(0x66,0x66,0x66))
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
r=p.add_run("최종 결과 보고서"); set_font(r,24,True,DARK)
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Unrailed! 모작 — 2인 로컬 대전 게임"); set_font(r,14,True,ACCENT)
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(10)
r=p.add_run("개발 인원: 2023184027 이현렬 · 2023184037 허찬휘     |     플랫폼: Win32 API + GDI+ (C++)")
set_font(r,10.5,False,RGBColor(0x44,0x44,0x44))

# ===== 1 =====
h1(1,"게임 소개 및 특징, 게임 규모")
h2("게임 소개")
body("본 프로젝트는 인기 협동 게임 'Unrailed'를 모작하여, 원작에 없던 PVP(1:1 대전) 요소를 더한 2인 로컬 멀티플레이 게임입니다. "
     "탑뷰 시점으로 한 화면을 위·아래로 나눠 두 플레이어가 동시에 플레이하며, 각자 자원을 채취해 선로를 깔고, 폭탄을 실은 기차를 "
     "상대의 기지에 먼저 도착시켜 기지를 폭파하는 것이 목표입니다. 원작이 협동에 초점을 둔 것과 달리, 본 게임은 장애물로 상대를 "
     "방해하고 먼저 폭탄을 배달하는 경쟁 구도로 재해석하였습니다.")
h2("게임의 특징")
bullet("기차는 멈추지 않고 계속 전진하므로, 선로가 끊기거나 장애물에 부딪히면 즉시 패배합니다.", bold_head="멈추지 않는 기차 — ")
bullet("맵의 나무·돌을 채집해 기지의 제작대에서 레일·장애물·폭탄을 제작합니다.", bold_head="자원 → 제작 루프 — ")
bullet("폭탄을 실은 기차로 상대 기지를 폭파하면 승리합니다.", bold_head="폭탄 배달 승부 — ")
bullet("상대 선로 앞에 장애물을 설치해 기차를 탈선시키는 방해 플레이가 가능합니다.", bold_head="방해 대전 — ")
bullet("한 화면을 위(P1)·아래(P2)로 분할하여 두 명이 한 키보드로 동시에 플레이합니다.", bold_head="로컬 2인 분할화면 — ")
h2("게임 규모")
bullet("타일 기반 맵은 60 × 30 타일이며, 타일당 32px로 실제 월드 크기는 1920 × 960px입니다.")
bullet("맵 양 끝에 각 플레이어의 기지와 제작대를 배치하였습니다.")
bullet("맵 곳곳에 나무·돌 자원을 배치하였고, 채집 후 일정 시간이 지나면 다시 생성됩니다.")
bullet("로컬 2인 멀티플레이로, 분할화면과 독립 카메라를 사용합니다.")
body("※ 제안서에서는 90×60 타일을 계획하였으나, 렌더링 성능과 게임 템포를 고려해 60×30으로 조정하였습니다. "
     "맵 크기 정의와 배경 이미지(1920×960)를 일치시켜 카메라 범위와 충돌 판정이 어긋나지 않도록 하였습니다.", size=9.5)

# ===== 2 =====
h1(2,"실행 방법 및 필요한 라이브러리")
h2("설치 및 실행 방법")
bullet("제출 폴더(Unrailed)에는 솔루션·소스·리소스·실행파일이 모두 포함되어 있습니다.")
bullet("간편 실행은 'Unrailed/Unrailed.exe'를 더블클릭하면 되며, 같은 폴더의 Image·Sound 폴더가 함께 있어야 이미지와 사운드가 정상 출력됩니다.")
bullet("소스 빌드는 'Unrailed.slnx'를 Visual Studio로 열고 구성을 'Release / x64'로 설정한 뒤 빌드·실행합니다.")
bullet("실행파일은 VC++ 런타임을 정적 링크(/MT)하여, 별도의 재배포 패키지 설치 없이 어떤 Windows PC에서도 실행됩니다.")
h2("조작 방법 (작동 방법)")
make_table(["구분","플레이어 1","플레이어 2"],
  [["화면","위쪽 화면","아래쪽 화면"],["이동","W A S D","방향키 ← ↑ ↓ →"],
   ["설치물 종류 변경","Q","1"],["레일 방향 전환","R","2"],["설치","E","3"],
   ["양동이 줍기 / 놓기","F","0"]], widths=[2.0,2.3,2.3])
make_table(["공통 / 시스템 키","기능"],
  [["ESC","일시정지 메뉴 (계속하기 / 설정 / 게임 종료)"],["F11","전체화면 ↔ 창모드 전환"],
   ["F2 / F3","무한 레일 모드 / 무한 자원 모드 (디버그용)"],
   ["마우스","시작화면 '게임시작'·'조작법' 버튼, 볼륨 슬라이더, 메뉴 버튼 클릭"]], widths=[2.2,4.4])
body("그 외 자동 동작으로, 플레이어가 자원에 접근하면 자동으로 채집되고, 물 양동이를 들고 기차에 닿으면 과열이 식습니다.", size=9.5)
h2("필요한 라이브러리 및 DLL")
bullet("Win32 API로 윈도우 생성과 메시지 루프, 입력을 처리합니다.")
bullet("GDI+(gdiplus.lib)로 비트맵·도형·텍스트를 렌더링합니다.")
bullet("WinMM(winmm.lib)으로 PlaySound 사운드 재생과 waveOutSetVolume 볼륨 조절을 수행합니다.")
bullet("ATL(atlimage.h의 CImage)로 플레이어 스프라이트 시트를 불러옵니다.")
body("위 라이브러리는 소스 내 #pragma comment(lib, ...)로 링크되며 모두 Windows에 기본 포함되어 있습니다. "
     "VC++ 런타임을 정적 링크하였으므로 별도로 제출해야 하는 DLL은 없습니다.", size=10)

# ===== 3 =====
h1(3,"구현한 내용")
h2("계획에 따라 구현된 내용")
bullet("플레이어 이동과 맵 이미지 색상 기반의 픽셀 단위 충돌·물 판정을 구현하였습니다.")
bullet("타일 맵 로드와 더블 버퍼링·독립 카메라 기반의 분할화면 렌더링을 구현하였습니다.")
bullet("자원의 자동 채집·드랍·리스폰과 기지 제작대에서의 레일·장애물·폭탄 제작을 구현하였습니다.")
bullet("레일은 상하좌우로 설치할 수 있으며, 이웃 레일을 인식해 곡선을 자동으로 생성하도록 구현하였습니다.")
bullet("기차는 레일을 따라 자동 주행하고, 과열되면 멈추며 물로 식힐 수 있도록 구현하였습니다.")
bullet("폭탄은 기차에 적재하거나 타일에 설치해 폭발시키며, 폭발 시 주변 장애물이 제거되도록 구현하였습니다.")
bullet("장애물을 통한 방해와 탈선·도착에 따른 승리/패배 판정을 구현하였습니다.")
bullet("시작화면·조작법·일시정지·설정(볼륨)·게임 시작 연출·승리 화면 등 UI를 구현하였습니다.")
bullet("배경음과 효과음, 볼륨 조절 기능을 구현하였습니다.")
h2("계획 대비 추가로 구현한 내용")
bullet("플레이어별 레일을 소유자로 분리하고 색상(P1 붉은색 / P2 푸른색)으로 구분하였습니다.")
bullet("기차의 진행 방향에 따른 부드러운 회전과 탈선 위기 경고 표시를 추가하였습니다.")
bullet("전체화면(F11)을 지원하며, 화면 확대 출력과 마우스 좌표 보정을 구현하였습니다.")
bullet("2D 그리드 기반 O(1) 조회와 회전·색상 이미지 사전 생성 등으로 성능을 최적화하였습니다.")
h2("계획했으나 구현하지 못했거나 변경한 내용 (사유 포함)")
bullet("맵 크기를 90×60에서 60×30으로 축소하였습니다. 분할화면을 매 프레임 2회 렌더링하는 구조상 큰 맵에서 성능 저하가 있었고, "
       "한 판의 진행 길이를 적절히 맞추기 위해 조정하였습니다.")
bullet("배경음을 MCI로 MP3 재생하려 하였으나 경로의 한글 문자 때문에 초기화에 실패하여, 음원을 WAV로 변환하고 PlaySound로 "
       "재생하도록 변경하였습니다.")
bullet("PlaySound가 한 번에 한 소리만 재생하는 한계 때문에 시작음·효과음·BGM을 겹쳐 재생하지 못하고 순차 재생으로 구성하였습니다.")

# ===== 4 =====
h1(4,"스크린샷 · 팀원 역할 · 개발 일정")
h2("스크린샷")
add_image(GS, 6.3, "[그림 1] 시작 화면 — 게임시작 버튼, 볼륨 슬라이더, 조작법(우상단) 버튼")
add_image(MAP, 6.3, "[그림 2] 게임 맵 — 양 끝 기지와 곳곳의 자원, 지형(땅/물)")
body("실제 플레이 화면(분할화면, 레일 설치, 기차 주행, 폭탄 폭발, 승리 화면 등) 캡처를 아래 영역에 추가로 삽입하면 됩니다.", size=9.5)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("[ 게임 플레이 스크린샷 삽입 영역 ]"); set_font(r,11,True,RGBColor(0x99,0x99,0x99))
doc.add_paragraph()
h2("팀원 역할 및 팀원 간 개발 내용")
make_table(["팀원","담당 개발 내용"],
  [["이현렬 (2023184027)","레일 설치 시스템, 폭탄 시스템, 열차 이동, 열차 과열 시스템, 승리/패배 판정"],
   ["허찬휘 (2023184037)","플레이어 이동, 맵 제작, 장애물 시스템, 자원 시스템"]], widths=[2.4,4.2])
body("두 팀원이 각자 담당한 시스템을 단일 소스(main.cpp)에서 통합하였으며, 공용 자료구조(맵·레일 그리드, 게임 상태)와 "
     "렌더링 파이프라인을 공유하여 기능을 연결하였습니다.", size=10)
h2("제작 일정에 따른 개발 상황")
make_table(["주차","계획","진행 상황"],
  [["1주차","플레이어 동작 및 맵, UI 제작","완료 — 이동·충돌·맵 로드·기본 UI 구현"],
   ["2주차","열차 시스템 제작","완료 — 레일 주행·과열·냉각 구현"],
   ["3주차","레일 시스템 제작","완료 — 상하좌우 설치·곡선 자동 연결"],
   ["4주차","자원 및 장애물 시스템","완료 — 채집·제작·장애물·폭탄"],
   ["5주차","테스트 및 버그 수정","완료 — 승패·사운드·전체화면 + 색 구분·회전·성능 최적화·버그 수정"]],
  widths=[1.2,2.7,2.9])

# ===== 5 =====
h1(5,"제작 후기")
body("이번 프로젝트를 통해 Win32 API와 GDI+만으로 실시간 게임을 구성하는 전체 흐름(메시지 루프, 더블 버퍼링, 분할화면 렌더링, "
     "입력 처리)을 직접 구현해 볼 수 있었습니다. 특히 레일이 많아질수록 발생하던 끊김 문제를 해결하는 과정에서, 단순히 그림을 "
     "그리는 것보다 '어떻게 적게·빠르게 그리느냐'가 중요하다는 것을 체감하였습니다. 매 프레임 반복되는 비용을 줄이기 위해 선형 탐색을 "
     "2D 그리드 기반 O(1) 조회로 바꾸고, 회전·색상을 입힌 레일 이미지를 미리 만들어 두는 등 자료구조와 사전 계산의 중요성을 "
     "배웠습니다. 또한 Debug와 Release 빌드의 성능 차이, 한글 경로로 인한 사운드 재생 문제처럼 환경에 따른 변수도 직접 부딪히며 "
     "해결하였습니다.")
body("레일의 곡선을 이웃 정보로 자동 판별하는 알고리즘과 두 플레이어의 레일을 소유자별로 분리하는 처리는 구현이 까다로웠지만, "
     "버그를 하나씩 잡아가며 게임이 점점 완성되어 가는 과정이 가장 큰 보람이었습니다. 협업 측면에서는 각자 맡은 시스템을 하나의 "
     "코드베이스로 통합하면서 모듈 간 인터페이스와 공용 상태 관리의 중요성을 느낄 수 있었습니다.")

# ===== 6 =====
h1(6,"유튜브 주소")
body("게임 실행 영상(시연) 주소는 다음과 같습니다.")
p=doc.add_paragraph()
r=p.add_run("▶ https://youtu.be/__________________   (업로드 후 주소 기입)"); set_font(r,11,True,RGBColor(0x10,0x40,0xA0))
body("※ 미등록(Unlisted)으로 업로드하면 검색에는 노출되지 않으면서 링크로 시청이 가능합니다.", size=9.5)

doc.save(OUT)
print("SAVED:", OUT)
